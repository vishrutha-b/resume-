import os
import io
import json
import docx
from unittest.mock import patch
from app import app

def create_production_grade_docx():
    doc = docx.Document()
    doc.add_heading('Alexander Chen - Lead Site Reliability Engineer', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Results-driven Lead SRE with 10+ years of experience in designing, implementing, and maintaining highly available, scalable, and secure infrastructure for global financial services. Expertise in Kubernetes configuration, cloud-native architectures (AWS/GCP), and establishing mature CI/CD pipelines.')
    
    doc.add_heading('Professional Experience', level=1)
    doc.add_heading('Lead Infrastructure Engineer | FinTech Global Solutions | New York, NY | Oct 2020 - Present', level=2)
    doc.add_paragraph('- Architected and migrating monolithic applications to a microservices architecture on AWS EKS, accommodating a 300% increase in trading volume and achieving 99.999% uptime.')
    doc.add_paragraph('- Built enterprise-wide CI/CD pipelines using GitLab CI, Jenkins, and ArgoCD, reducing mean time to deployment across 50+ engineering teams from weeks to under 4 hours.')
    doc.add_paragraph('- Implemented Infrastructure as Code (IaC) governance using Terraform and Terragrunt, standardizing cloud provisioning securely and automatically tracking infrastructure drift.')
    doc.add_paragraph('- Integrated comprehensive observability tooling with Prometheus, Grafana, and ELK stack, decreasing MTTR during P1 incidents by 45%.')
    
    doc.add_heading('Senior DevOps Engineer | Capital Markets Inc. | Chicago, IL | Jan 2016 - Sept 2020', level=2)
    doc.add_paragraph('- Managed highly regulated, on-premise infrastructure consisting of 5,000+ Linux instances (RHEL/CentOS) configuring via Ansible and Puppet.')
    doc.add_paragraph('- Established automated security compliance scanning in the pipeline using SonarQube, Trivy, and HashiCorp Vault for secrets management.')
    doc.add_paragraph('- Migrated legacy Oracle databases to PostgreSQL clusters safely with minimal downtime using logical replication strategies.')
    
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph('Cloud & Containerization: AWS (EC2, EKS, RDS, S3, IAM), Docker, Kubernetes, Helm')
    doc.add_paragraph('CI/CD & Version Control: Jenkins, GitLab CI/CD, ArgoCD, GitHub Actions, Git')
    doc.add_paragraph('Infrastructure & Config Gov: Terraform, Ansible, Chef, Puppet, Packer')
    doc.add_paragraph('Monitoring & Observability: Prometheus, Grafana, ELK Stack, Datadog, Splunk')
    doc.add_paragraph('Languages: Python (Boto3, Flask), Bash, Go')
    doc.add_paragraph('Security: HashiCorp Vault, Trivy, IAM, OPA (Open Policy Agent)')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. in Computer Science | University of Illinois at Urbana-Champaign (UIUC) | 2015')
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def get_production_job_description():
    return """
    Title: VP, Lead DevOps & Cloud Platform Engineer
    Department: Global Markets Technology
    Location: New York, NY
    
    Role Overview:
    We are seeking a VP-level Lead DevOps and infrastructure specialist to govern our automated trading platform infrastructure. The ideal candidate will bridge the gap between development and operations by strictly enforcing IaC methodologies, hardening cloud environments, and driving Kubernetes adoption securely.
    
    Mandatory Requirements:
    - 8+ years in SRE, DevOps, or Infrastructure roles, preferably within highly regulated environments (FinServ/Healthcare).
    - Deep expertise in Kubernetes (EKS/GKE) administration, networking, and security profiles.
    - Master-level understanding of Infrastructure as Code using Terraform and configuration management with Ansible.
    - Extensive experience managing extensive CI/CD lifecycles using Jenkins and ArgoCD or Spinnaker.
    - Strong programming proficiency in Python or Go for infrastructure automation.
    - Production experience with enterprise observability suites (Datadog, Splunk, Prometheus/Grafana).
    
    Good-to-Haves:
    - Experience with Zero Trust network architectures or HashiCorp Vault.
    - Familiarity with trading or fixing protocols.
    - Cloud certifications (e.g., AWS Certified Solutions Architect/DevOps Professional).
    """

def get_mocked_llm_response():
    # Production-grade evaluation JSON mapping exactly to ATSResponseSchema
    return json.dumps({
        "ats_score": 92,
        "skill_match": {
            "matched": ["Kubernetes", "AWS EKS", "Terraform", "Ansible", "Jenkins", "ArgoCD", "Python", "Prometheus", "Grafana", "HashiCorp Vault"],
            "missing": ["Trading/FIX Protocols", "Spinnaker", "GKE"],
            "partial": ["Datadog", "Splunk"]
        },
        "domain_fit": "High",
        "experience_fit": "High",
        "key_gaps": [
            "Lack of explicit experience with Trading/FIX protocols, which was listed as a 'good-to-have' for the Global Markets Technology department.",
            "Candidate mentioned ELK but explicitly missing production Datadog usage, which is preferred."
        ],
        "strong_points": [
            "Exceptional alignment with mandatory Kubernetes, EKS, and Terraform requirements.",
            "Demonstrated enterprise-scale impact (migrating monoliths to microservices, 300% volume increase).",
            "Strong background in financial services ('Capital Markets Inc.', 'FinTech Global Solutions'), satisfying the highly regulated environment prerequisite.",
            "Relevant proficiency in Python and advanced CI/CD tooling (Jenkins, ArgoCD)."
        ],
        "recommendations": [
            "Proceed to technical interview prioritizing questions on EKS networking and security profiles.",
            "Assess candidate's ability to quickly adapt to Datadog/Splunk based on their strong Prometheus/ELK experience.",
            "Inquire further about practical applications of HashiCorp Vault within their CI/CD pipelines."
        ]
    })

class MockResponse:
    def __init__(self, json_data, status_code):
        self.json_data = json_data
        self.status_code = status_code

    def json(self):
        return self.json_data

    def raise_for_status(self):
        if self.status_code != 200:
            raise Exception("HTTP Error mocked")

@patch('requests.post')
def test_screen_resume(mock_post):
    # Setup mock
    mock_response_data = {
        "choices": [
            {
                "message": {
                    "content": get_mocked_llm_response()
                }
            }
        ]
    }
    mock_post.return_value = MockResponse(mock_response_data, 200)

    # Initialize Flask test client
    client = app.test_client()
    
    # Prepare form data
    data = {
        'job_desc': get_production_job_description(),
    }
    
    # Simulate file upload
    docx_stream = create_production_grade_docx()
    data['resume'] = (docx_stream, 'alexander_chen_resume.docx')
    
    print("Sending request to /api/screen_resume endpoint using mocked LLM...")
    response = client.post('/api/screen_resume', data=data, content_type='multipart/form-data')
    
    if response.status_code != 201:
        print(f"FAILED: Server returned status {response.status_code}")
        print(response.data.decode('utf-8'))
        return
    
    if response.is_json:
        result_json = response.get_json()
        with open('output_mocked.json', 'w', encoding='utf-8') as f:
            json.dump(result_json, f, indent=2)
        print("Success: Wrote production-grade response to output_mocked.json")
    else:
        print(response.data.decode('utf-8'))

if __name__ == '__main__':
    test_screen_resume()
