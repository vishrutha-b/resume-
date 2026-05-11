import docx
doc = docx.Document()
doc.add_heading('Testing Animations Candidate', 0)
doc.add_paragraph('Exceptional Senior Fullstack Software Engineer.')
doc.add_paragraph('Skills: React, Next.js, WebGL, Python, Flask, Postgres, AWS, CI/CD pipelines.')
doc.add_paragraph('Built incredibly premium dashboard interfaces and handled robust ATS backend evaluations.')
doc.save('dummy_resume.docx')
print('Created dummy_resume.docx')
